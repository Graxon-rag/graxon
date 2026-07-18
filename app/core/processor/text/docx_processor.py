from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument
from app.utils.logger import logger
from .processor import Processor
from app.config.env import Env
from typing import List, Tuple


class DOCXProcessor(Processor):
    def __init__(self,
        file_path: str,
        filename: str,
        chunk_number: int,
        rag_chunk_start_index: int,
        pages_per_batch: int = 20,          # treated as paragraphs_per_batch for DOCX
        rag_chunk_size: int = Env.CHUNK_SIZE,
        rag_chunk_overlap: int = Env.CHUNK_OVERLAP,
        tail_carry_chars: int = 500,
    ):
        self.file_path = file_path
        self.filename = filename
        self.chunk_number = chunk_number
        self.rag_chunk_start_index = rag_chunk_start_index
        self.pages_per_batch = pages_per_batch
        self.rag_chunk_size = rag_chunk_size
        self.rag_chunk_overlap = rag_chunk_overlap
        self.tail_carry_chars = tail_carry_chars

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=rag_chunk_size,
            chunk_overlap=rag_chunk_overlap,
            separators=["\n\n", "\n", ".", " ", ""],
        )

    async def process(self) -> Tuple[List[Document], int, bool]:
        """
        DOCX has no true page concept — paragraphs are the natural unit.
        pages_per_batch maps to paragraphs_per_batch here.
        Extracts a batch of paragraphs, joins as plain text,
        then splits with RecursiveCharacterTextSplitter.

        Returns:
            documents: list of Document
            next_rag_chunk_start_index: pass this to next queue message
            is_last: True if this was the final batch
        """
        try:
            docx = DocxDocument(self.file_path)

            # filter out empty paragraphs
            paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
            total_paragraphs = len(paragraphs)

            para_start = self.chunk_number * self.pages_per_batch
            if para_start >= total_paragraphs:
                raise ValueError(
                    f"chunk_number {self.chunk_number} is out of range. "
                    f"Total paragraphs: {total_paragraphs}, pages_per_batch: {self.pages_per_batch}"
                )

            para_end = min(para_start + self.pages_per_batch, total_paragraphs)
            is_last = para_end >= total_paragraphs

            # carry tail from previous batch to patch boundary cuts
            tail_text = ""
            if para_start > 0:
                prev_batch_text = "\n".join(paragraphs[max(0, para_start - 3):para_start])
                tail_text = prev_batch_text[-self.tail_carry_chars:]

            batch_text = "\n".join(paragraphs[para_start:para_end])
            raw_text = tail_text + batch_text

            documents = self._split_into_rag_chunks(raw_text, para_start)
            return documents, self.rag_chunk_start_index + len(documents), is_last

        except Exception as e:
            logger.error(f"Failed to process DOCX file {self.file_path}. Error: {e}")
            raise e

    def _split_into_rag_chunks(self, raw_text: str, para_start: int) -> List[Document]:
        texts = self.splitter.split_text(raw_text)

        documents = []
        for i, text in enumerate(texts):
            absolute_index = self.rag_chunk_start_index + i
            doc = Document(
                id=f"{self.filename}-{absolute_index}",
                page_content=text,
                metadata={
                    "source": self.file_path,
                    "file_chunk_number": self.chunk_number,
                    "rag_chunk_number": absolute_index,
                    "paragraph_start": para_start,      # which paragraph batch this came from
                },
            )
            documents.append(doc)

        return documents
